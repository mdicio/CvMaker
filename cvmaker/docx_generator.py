"""DOCX generator — creates Word documents from parsed CV data."""

import json
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from .parser import CV, Section, SubSection, Paragraph, ListItem, TextRun


# ── Colour helpers ──────────────────────────────────────────────────────────


def hex_to_rgb(hex_color: str) -> RGBColor:
    """Convert hex colour string to *RGBColor*."""
    hex_color = hex_color.lstrip("#")
    r, g, b = (int(hex_color[i : i + 2], 16) for i in (0, 2, 4))
    return RGBColor(r, g, b)


def load_template(template_path: str | None = None) -> dict:
    """Load a JSON styling template."""
    if template_path is None:
        template_path = Path(__file__).parent / "templates" / "default.json"
    with open(template_path, "r") as f:
        return json.load(f)


# ── Hyperlink helper ────────────────────────────────────────────────────────


def _add_hyperlink(
    paragraph, text: str, url: str, font_name: str, font_size: float, color: RGBColor
) -> None:
    """Insert a clickable hyperlink run into *paragraph*."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_el = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)
    rPr.append(rFonts)

    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(font_size * 2)))
    rPr.append(sz)

    clr = OxmlElement("w:color")
    clr.set(qn("w:val"), str(color))
    rPr.append(clr)

    bold_el = OxmlElement("w:b")
    rPr.append(bold_el)

    run_el.append(rPr)

    t_el = OxmlElement("w:t")
    t_el.set(qn("xml:space"), "preserve")
    t_el.text = text
    run_el.append(t_el)

    hyperlink.append(run_el)
    paragraph._element.append(hyperlink)


# ── Run formatting ──────────────────────────────────────────────────────────


def add_formatted_runs(
    paragraph,
    runs: list[TextRun],
    font_name: str,
    font_size: int,
    body_color: RGBColor,
    accent_color: RGBColor | None = None,
) -> None:
    """Add formatted text runs (with optional hyperlinks) to a paragraph."""
    for rd in runs:
        if rd.url and accent_color:
            _add_hyperlink(
                paragraph, rd.text, rd.url, font_name, font_size, accent_color
            )
        else:
            run = paragraph.add_run(rd.text)
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.color.rgb = body_color
            run.bold = rd.bold
            run.italic = rd.italic


# ── Cell helpers ────────────────────────────────────────────────────────────


def _shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tc_pr.append(shd)


def _set_cell_margins(cell, pad_pts: float) -> None:
    mar = OxmlElement("w:tcMar")
    pad_twips = str(int(Pt(pad_pts)))
    for side in ("top", "bottom", "start", "end"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), pad_twips)
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    cell._tc.get_or_add_tcPr().append(mar)


def _remove_cell_borders(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_el = OxmlElement(f"w:{edge}")
        edge_el.set(qn("w:val"), "nil")
        borders.append(edge_el)
    tc_pr.append(borders)


# ── Document generation ─────────────────────────────────────────────────────


def generate_docx(cv: CV, output_path: str, template_path: str | None = None) -> str:
    """Generate a DOCX file from parsed CV data."""
    template = load_template(template_path)
    doc = Document()

    doc_section = doc.sections[0]

    def _usable_width():
        return (
            doc_section.page_width - doc_section.left_margin - doc_section.right_margin
        )

    for sec in doc.sections:
        sec.top_margin = Inches(template["margins"]["top"])
        sec.bottom_margin = Inches(template["margins"]["bottom"])
        sec.left_margin = Inches(template["margins"]["left"])
        sec.right_margin = Inches(template["margins"]["right"])

    fonts = template["fonts"]
    title_size = fonts.get(
        "title_size", fonts.get("heading2_size", fonts["heading1_size"] - 2)
    )
    colors = template["colors"]
    spacing = template["spacing"]
    header_cfg = template.get("header", {})
    accent_rgb = hex_to_rgb(colors.get("accent", colors["body"]))

    # ── Header block ────────────────────────────────────────────────────
    if cv.name or cv.contact:
        rows = 1 + (1 if cv.contact else 0)
        header_table = doc.add_table(rows=rows, cols=1)
        header_table.autofit = False
        header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        width = _usable_width()
        header_table.columns[0].width = width

        for cell in header_table._cells:
            _shade_cell(
                cell, header_cfg.get("background", colors.get("header_bg", "eaf3ff"))
            )
            _remove_cell_borders(cell)
            cell.width = width
            _set_cell_margins(cell, header_cfg.get("padding", spacing["section_after"]))

        # Name row
        name_cell = header_table.rows[0].cells[0]
        name_para = name_cell.paragraphs[0]
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pad = header_cfg.get("padding", spacing["section_after"])
        name_para.paragraph_format.space_after = Pt(pad * 0.25)
        name_para.paragraph_format.space_before = Pt(pad * 0.15)

        name_run = name_para.add_run(cv.name)
        name_run.font.name = fonts["name"]
        name_run.font.size = Pt(fonts["heading1_size"])
        name_run.font.color.rgb = hex_to_rgb(colors["heading"])
        name_run.bold = True

        if cv.subtitle:
            sep_run = name_para.add_run(" | ")
            sep_run.font.name = fonts["name"]
            sep_run.font.size = Pt(title_size)
            sep_run.font.color.rgb = hex_to_rgb(colors["body"])

            title_run = name_para.add_run(cv.subtitle)
            title_run.font.name = fonts["name"]
            title_run.font.size = Pt(title_size)
            title_run.font.color.rgb = hex_to_rgb(colors["heading"])
            title_run.bold = True

        # Contact row
        if cv.contact:
            contact_cell = header_table.rows[1].cells[0]
            contact_para = contact_cell.paragraphs[0]
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            contact_para.paragraph_format.space_after = Pt(
                header_cfg.get("gap", spacing["section_after"])
            )
            contact_para.paragraph_format.line_spacing = 1.0
            add_formatted_runs(
                contact_para,
                cv.contact.runs,
                fonts["name"],
                fonts["body_size"],
                hex_to_rgb(colors["body"]),
                accent_rgb,
            )

    # ── Sections ────────────────────────────────────────────────────────
    for section in cv.sections:
        _add_section(doc, section, fonts, colors, spacing, accent_rgb)

    doc.save(output_path)
    return output_path


def _add_section(doc, section: Section, fonts, colors, spacing, accent_rgb) -> None:
    """Add a section to the document."""
    header_para = doc.add_paragraph()
    header_para.paragraph_format.space_before = Pt(spacing["section_before"])
    header_para.paragraph_format.space_after = Pt(0)
    header_run = header_para.add_run(section.title)
    header_run.font.name = fonts["name"]
    header_run.font.size = Pt(fonts["heading2_size"])
    header_run.font.color.rgb = hex_to_rgb(colors["heading"])
    header_run.bold = True

    border_para = doc.add_paragraph()
    border_para.paragraph_format.space_before = Pt(0)
    border_para.paragraph_format.space_after = Pt(spacing["section_after"])
    border_run = border_para.add_run("_" * 80)
    border_run.font.size = Pt(2)
    border_run.font.color.rgb = hex_to_rgb(colors["accent"])

    for item in section.content:
        if isinstance(item, SubSection):
            _add_subsection(doc, item, fonts, colors, spacing, accent_rgb)
        elif isinstance(item, Paragraph):
            _add_paragraph(doc, item, fonts, colors, spacing, accent_rgb)
        elif isinstance(item, ListItem):
            _add_list_item(doc, item, fonts, colors, spacing, accent_rgb)


def _add_subsection(
    doc, subsection: SubSection, fonts, colors, spacing, accent_rgb
) -> None:
    """Add a subsection to the document."""
    doc_section = doc.sections[0]
    usable_width = (
        doc_section.page_width - doc_section.left_margin - doc_section.right_margin
    )

    header_para = doc.add_paragraph()
    header_para.paragraph_format.space_before = Pt(spacing["section_before"] // 2)
    header_para.paragraph_format.space_after = Pt(spacing["section_after"] // 2)
    header_para.paragraph_format.tab_stops.add_tab_stop(
        usable_width, WD_TAB_ALIGNMENT.RIGHT
    )

    if subsection.url:
        _add_hyperlink(
            header_para,
            subsection.title,
            subsection.url,
            fonts["name"],
            fonts["heading3_size"],
            hex_to_rgb(colors.get("accent", colors["subheading"])),
        )
    else:
        header_run = header_para.add_run(subsection.title)
        header_run.font.name = fonts["name"]
        header_run.font.size = Pt(fonts["heading3_size"])
        header_run.font.color.rgb = hex_to_rgb(colors["subheading"])
        header_run.bold = True

    if subsection.date:
        header_para.add_run().add_tab()
        date_run = header_para.add_run(subsection.date)
        date_run.font.name = fonts["name"]
        date_run.font.size = Pt(fonts["body_size"])
        date_run.font.color.rgb = hex_to_rgb(colors["body"])
        date_run.italic = True

    for item in subsection.content:
        if isinstance(item, Paragraph):
            _add_paragraph(doc, item, fonts, colors, spacing, accent_rgb)
        elif isinstance(item, ListItem):
            _add_list_item(doc, item, fonts, colors, spacing, accent_rgb)


def _add_paragraph(doc, para: Paragraph, fonts, colors, spacing, accent_rgb) -> None:
    """Add a paragraph to the document."""
    doc_section = doc.sections[0]
    usable_width = (
        doc_section.page_width - doc_section.left_margin - doc_section.right_margin
    )

    doc_para = doc.add_paragraph()
    doc_para.paragraph_format.space_after = Pt(spacing["section_after"] // 2)

    if para.date:
        doc_para.paragraph_format.tab_stops.add_tab_stop(
            usable_width, WD_TAB_ALIGNMENT.RIGHT
        )

    add_formatted_runs(
        doc_para,
        para.runs,
        fonts["name"],
        fonts["body_size"],
        hex_to_rgb(colors["body"]),
        accent_rgb,
    )

    if para.date:
        doc_para.add_run().add_tab()
        date_run = doc_para.add_run(para.date)
        date_run.font.name = fonts["name"]
        date_run.font.size = Pt(fonts["body_size"])
        date_run.font.color.rgb = hex_to_rgb(colors["body"])
        date_run.italic = True


def _add_list_item(doc, item: ListItem, fonts, colors, spacing, accent_rgb) -> None:
    """Add a list item to the document."""
    doc_section = doc.sections[0]
    usable_width = (
        doc_section.page_width - doc_section.left_margin - doc_section.right_margin
    )

    doc_para = doc.add_paragraph(style="List Bullet")
    doc_para.paragraph_format.space_after = Pt(0)
    doc_para.paragraph_format.left_indent = Inches(0.25)

    if item.date:
        doc_para.paragraph_format.tab_stops.add_tab_stop(
            usable_width, WD_TAB_ALIGNMENT.RIGHT
        )

    add_formatted_runs(
        doc_para,
        item.runs,
        fonts["name"],
        fonts["body_size"],
        hex_to_rgb(colors["body"]),
        accent_rgb,
    )

    if item.date:
        doc_para.add_run().add_tab()
        date_run = doc_para.add_run(item.date)
        date_run.font.name = fonts["name"]
        date_run.font.size = Pt(fonts["body_size"])
        date_run.font.color.rgb = hex_to_rgb(colors["body"])
        date_run.italic = True
